from docx import Document
from datetime import date

def generate_agreement(input_path, output_path, values):
    doc = Document(input_path)

    for para in doc.paragraphs:
        for key, val in values.items():
            if f"{{{{{key}}}}}" in para.text:
                para.text = para.text.replace(f"{{{{{key}}}}}", val)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, val in values.items():
                    if f"{{{{{key}}}}}" in cell.text:
                        cell.text = cell.text.replace(f"{{{{{key}}}}}", val)

    doc.save(output_path)
    print(f"\n✅ Agreement saved as: {output_path}")


def main():
    print("\n📄 Welcome to Sara's Agreement Generator\n")

    # Ask user for inputs
    brand_name = input("Brand Name: ")
    company_name = input("Company Legal Name: ")
    company_address = input("Company Address: ")
    industry = input("Industry: ")
    flat_fee = input("Flat Fee (e.g., 1000): ")
    deposit = input("Deposit Amount (e.g., 10000): ")
    deposit_in_words = input("Deposit in Words (e.g., Ten Thousand Only): ")
    start_date = input(f"Agreement Date [default: today {date.today()}]: ") or str(date.today())

    values = {
        "brand_name": brand_name,
        "company_name": company_name,
        "company_address": company_address,
        "industry": industry,
        "flat_fee": flat_fee,
        "deposit": deposit,
        "deposit_in_words": deposit_in_words,
        "start_date": start_date
    }

    input_template = "Partnership Agreement Template.docx"
    output_filename = f"{brand_name.replace(' ', '_')}_Agreement.docx"
    generate_agreement(input_template, output_filename, values)


if __name__ == "__main__":
    main()
