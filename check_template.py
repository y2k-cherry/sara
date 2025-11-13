#!/usr/bin/env python3
"""Check what placeholders are in the Advance Deposit Invoice Template"""

from docx import Document

TEMPLATE_PATH = "Advance Deposit Invoice Template.docx"

def check_template():
    doc = Document(TEMPLATE_PATH)
    
    print("=" * 60)
    print("CHECKING TEMPLATE PLACEHOLDERS")
    print("=" * 60)
    
    placeholders_found = set()
    
    # Check paragraphs
    print("\n📄 PARAGRAPHS:")
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            print(f"  {i}: {p.text}")
            # Find placeholders
            import re
            matches = re.findall(r'\{\{([^}]+)\}\}', p.text)
            for match in matches:
                placeholders_found.add(match)
    
    # Check tables
    print("\n📊 TABLES:")
    for table_idx, table in enumerate(doc.tables):
        print(f"\n  Table {table_idx}:")
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                if cell.text.strip():
                    print(f"    Row {row_idx}, Cell {cell_idx}: {cell.text}")
                    # Find placeholders
                    import re
                    matches = re.findall(r'\{\{([^}]+)\}\}', cell.text)
                    for match in matches:
                        placeholders_found.add(match)
    
    print("\n" + "=" * 60)
    print("PLACEHOLDERS FOUND:")
    print("=" * 60)
    for p in sorted(placeholders_found):
        print(f"  • {p}")
    
    print("\n" + "=" * 60)
    print("EXPECTED PLACEHOLDERS IN CODE:")
    print("=" * 60)
    expected = [
        "Invoice_Number",
        "Brand_Name",
        "Brand_Address_Line_1",
        "Brand_Address_Line_2",
        "City",
        "State",
        "Pin_Code",
        "Phone",
        "Email",
        "Invoice_Date",
        "Due_Date",
        "Amount_Due",
        "Deposit_Amount",
        "Sub_Total"
    ]
    for e in expected:
        status = "✅" if e in placeholders_found else "❌"
        print(f"  {status} {e}")
    
    # Check for special Invoice_Number format
    print("\n" + "=" * 60)
    print("CHECKING FOR SPECIAL FORMATS:")
    print("=" * 60)
    
    # Re-check for #{{Invoice_Number}}
    full_text = "\n".join([p.text for p in doc.paragraphs])
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text += "\n" + cell.text
    
    if "#{{Invoice_Number}}" in full_text:
        print("  ✅ Found: #{{Invoice_Number}}")
    elif "{{Invoice_Number}}" in full_text:
        print("  ⚠️  Found: {{Invoice_Number}} (without #)")
    else:
        print("  ❌ Invoice_Number not found in any format")

if __name__ == "__main__":
    check_template()
