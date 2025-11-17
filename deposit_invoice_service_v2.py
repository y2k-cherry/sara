#!/usr/bin/env python3
"""
Deposit Invoice Generator Service - V2 with Enhanced Logging and Validation
Generates deposit invoices using the Advance Deposit Invoice Template
"""

import os
import re
from datetime import date, timedelta
from docx import Document
from slack_sdk import WebClient
from dotenv import load_dotenv
import json
from typing import Dict, List, Tuple, Optional

# Google Docs-based PDF exporter
from google_pdf import convert_docx_to_pdf_google as convert_docx_to_pdf

load_dotenv()

# Constants
SLACK_TOKEN = os.getenv("SLACK_BOT_TOKEN")
TEMPLATE_PATH = "Advance Deposit Invoice Template.docx"

# Initialize Slack client
slack_client = WebClient(token=SLACK_TOKEN)

# Required fields for invoice generation
REQUIRED_FIELDS = [
    "brand_name",
    "deposit_amount",
    "invoice_number"
]

# State management for deposit invoice generation flow
# Tracks threads that are in the middle of invoice generation
deposit_invoice_threads = {}  # thread_id -> {"stage": "awaiting_amount|awaiting_invoice_number", "brand_data": {...}, "amount": "..."}


def is_in_deposit_invoice_flow(thread_id: str) -> bool:
    """Check if a thread is currently in a deposit invoice generation flow"""
    return thread_id in deposit_invoice_threads


def get_deposit_invoice_state(thread_id: str) -> Optional[dict]:
    """Get the current state of a deposit invoice flow"""
    return deposit_invoice_threads.get(thread_id)


def set_deposit_invoice_state(thread_id: str, stage: str, brand_data: Optional[dict] = None, amount: str = ""):
    """Set the state for a deposit invoice flow"""
    deposit_invoice_threads[thread_id] = {
        "stage": stage,
        "brand_data": brand_data,
        "amount": amount
    }
    print(f"📝 [INVOICE STATE] Thread {thread_id}: stage={stage}, has_brand_data={bool(brand_data)}, amount={amount}")


def clear_deposit_invoice_state(thread_id: str):
    """Clear the state for a deposit invoice flow"""
    if thread_id in deposit_invoice_threads:
        print(f"🧹 [INVOICE STATE] Clearing state for thread {thread_id}")
        del deposit_invoice_threads[thread_id]


class InvoiceLogger:
    """Enhanced logging for invoice generation with stage tracking"""
    
    def __init__(self, thread_ts: str):
        self.thread_ts = thread_ts
        self.logs = []
        self.stage = "INITIALIZATION"
    
    def log(self, message: str, level: str = "INFO"):
        """Log a message with timestamp and stage"""
        log_entry = {
            "thread_ts": self.thread_ts,
            "stage": self.stage,
            "level": level,
            "message": message
        }
        self.logs.append(log_entry)
        
        # Also print to console for debugging
        emoji = {
            "INFO": "ℹ️",
            "SUCCESS": "✅",
            "WARNING": "⚠️",
            "ERROR": "❌",
            "DEBUG": "🔍"
        }.get(level, "📝")
        
        print(f"{emoji} [{self.stage}] {message}")
    
    def set_stage(self, stage: str):
        """Update current processing stage"""
        self.stage = stage
        self.log(f"Entering stage: {stage}", "DEBUG")
    
    def get_summary(self) -> str:
        """Get a formatted summary of all logs"""
        summary = f"📊 Invoice Generation Log Summary (Thread: {self.thread_ts})\n"
        summary += "=" * 60 + "\n\n"
        
        for log in self.logs:
            summary += f"[{log['stage']}] {log['level']}: {log['message']}\n"
        
        return summary


def clean_text(text: str) -> str:
    """Strip Slack mentions and extra whitespace."""
    cleaned = re.sub(r"<@[\w]+>", "", text).strip()
    return cleaned


def format_currency(value: str) -> str:
    """Format an integer/float string as Indian-rupee currency."""
    try:
        num = float(value)
        formatted = f"₹{num:,.0f}"
        return formatted
    except Exception as e:
        print(f"⚠️ Error formatting currency '{value}': {e}")
        return value


def convert_number_to_words(number_str: str) -> str:
    """Convert a number string to words (Indian numbering system)."""
    try:
        num = int(float(number_str))
        
        if num == 0:
            return "zero"
        elif num < 0:
            return "negative " + convert_number_to_words(str(-num))
        
        # Indian numbering system
        ones = ["", "one", "two", "three", "four", "five", "six", "seven", "eight", "nine"]
        teens = ["ten", "eleven", "twelve", "thirteen", "fourteen", "fifteen", 
                "sixteen", "seventeen", "eighteen", "nineteen"]
        tens = ["", "", "twenty", "thirty", "forty", "fifty", "sixty", "seventy", "eighty", "ninety"]
        
        def convert_hundreds(n):
            result = ""
            if n >= 100:
                result += ones[n // 100] + " hundred "
                n %= 100
            if n >= 20:
                result += tens[n // 10] + " "
                n %= 10
            elif n >= 10:
                result += teens[n - 10] + " "
                n = 0
            if n > 0:
                result += ones[n] + " "
            return result.strip()
        
        result = ""
        if num >= 10000000:  # 1 crore
            crores = num // 10000000
            result = convert_hundreds(crores) + " crore "
            num %= 10000000
            
        if num >= 100000:  # 1 lakh
            lakhs = num // 100000
            result += convert_hundreds(lakhs) + " lakh "
            num %= 100000
            
        if num >= 1000:  # 1 thousand
            thousands = num // 1000
            result += convert_hundreds(thousands) + " thousand "
            num %= 1000
            
        if num > 0:
            result += convert_hundreds(num)
            
        return result.strip()
        
    except Exception as e:
        print(f"⚠️ Error converting number to words: {e}")
        return number_str


def extract_deposit_amount(message_text: str, logger: InvoiceLogger) -> str:
    """
    Extract deposit amount from the message.
    Patterns: "5000", "Rs 5000", "₹5000", "deposit 5000", "amount 5000" etc.
    """
    logger.set_stage("EXTRACT_AMOUNT")
    logger.log(f"Attempting to extract amount from: '{message_text[:100]}...'", "DEBUG")
    
    # Try various patterns - more specific patterns first
    patterns = [
        (r'(?:amount|deposit)\s+(?:of\s+)?(?:rs\.?|₹)?\s*([0-9,]+)', "Pattern: 'amount/deposit of Rs X'"),
        (r'(?:for|invoice\s+for)\s+(?:rs\.?|₹)?\s*([0-9,]+)', "Pattern: 'for Rs X'"),
        (r'(?:rs\.?|₹)\s*([0-9,]+)', "Pattern: 'Rs X' or '₹X'"),
        (r'\b([0-9]{4,})\b', "Pattern: 'Any 4+ digit number'"),
    ]
    
    for pattern, description in patterns:
        logger.log(f"Trying {description}", "DEBUG")
        match = re.search(pattern, message_text, re.IGNORECASE)
        if match:
            amount = match.group(1).replace(',', '')
            logger.log(f"Successfully extracted amount: {amount} using {description}", "SUCCESS")
            return amount
    
    logger.log("No deposit amount found in message", "WARNING")
    return ""


def extract_invoice_number(message_text: str, logger: InvoiceLogger) -> str:
    """
    Extract invoice number from the message.
    Patterns: "invoice #123", "invoice number 123", "#INV-001", etc.
    """
    logger.set_stage("EXTRACT_INVOICE_NUMBER")
    logger.log(f"Attempting to extract invoice number from: '{message_text[:100]}...'", "DEBUG")
    
    patterns = [
        (r'#\s*([A-Z0-9/-]+)', "Pattern: '#INV-001' or '#123'"),
        (r'invoice\s+#\s*([A-Z0-9/-]+)', "Pattern: 'invoice #123'"),
        (r'invoice\s+number\s*:?\s*([A-Z0-9/-]+)', "Pattern: 'invoice number 123'"),
        (r'\b(INV-[0-9]+)\b', "Pattern: 'INV-001'"),
        (r'\b([A-Z]{2,3}/[A-Z]{2}/[0-9]+)\b', "Pattern: 'SB/DP/001'"),
    ]
    
    for pattern, description in patterns:
        logger.log(f"Trying {description}", "DEBUG")
        match = re.search(pattern, message_text, re.IGNORECASE)
        if match:
            invoice_num = match.group(1).strip().upper()
            logger.log(f"Successfully extracted invoice number: {invoice_num} using {description}", "SUCCESS")
            return invoice_num
    
    logger.log("No invoice number found in message", "WARNING")
    return ""


def parse_address_components(address: str, logger: InvoiceLogger) -> dict:
    """
    Parse a combined address into separate components.
    Returns dict with Brand_Address_Line_1, Brand_Address_Line_2, City, State, Pin_Code
    """
    logger.set_stage("PARSE_ADDRESS")
    logger.log(f"Parsing address: '{address}'", "DEBUG")
    
    if not address:
        logger.log("No address provided, returning empty components", "WARNING")
        return {
            "Brand_Address_Line_1": "",
            "Brand_Address_Line_2": "",
            "City": "",
            "State": "",
            "Pin_Code": ""
        }
    
    # Split address by commas
    parts = [p.strip() for p in address.split(',')]
    logger.log(f"Address split into {len(parts)} parts: {parts}", "DEBUG")
    
    # Initialize components
    components = {
        "Brand_Address_Line_1": "",
        "Brand_Address_Line_2": "",
        "City": "",
        "State": "",
        "Pin_Code": ""
    }
    
    # Try to extract pin code (6 digits)
    pin_code_pattern = r'\b(\d{6})\b'
    for i, part in enumerate(parts):
        match = re.search(pin_code_pattern, part)
        if match:
            components["Pin_Code"] = match.group(1)
            logger.log(f"Found pin code: {components['Pin_Code']}", "SUCCESS")
            # Remove pin code from this part
            parts[i] = re.sub(pin_code_pattern, '', part).strip()
            break
    
    # Remove empty parts
    parts = [p for p in parts if p]
    
    # Assign remaining parts
    if len(parts) >= 1:
        components["Brand_Address_Line_1"] = parts[0]
    if len(parts) >= 2:
        components["Brand_Address_Line_2"] = parts[1]
    if len(parts) >= 3:
        components["City"] = parts[2]
    if len(parts) >= 4:
        components["State"] = parts[3]
    
    # If we have fewer parts, try to be smart about city/state
    if len(parts) == 3 and not components["State"]:
        components["State"] = parts[2]
        components["City"] = parts[1] if len(parts) > 1 else ""
    
    logger.log(f"Parsed address components: {json.dumps(components, indent=2)}", "SUCCESS")
    return components


def extract_invoice_fields(message_text: str, brand_data: Optional[dict], logger: InvoiceLogger) -> Tuple[dict, List[str]]:
    """
    Extract invoice fields from message and brand data.
    
    Args:
        message_text: The user's message
        brand_data: Optional brand data from brand_info_service
        logger: InvoiceLogger instance
    
    Returns:
        Tuple of (values dict, list of missing fields)
    """
    logger.set_stage("EXTRACT_FIELDS")
    logger.log("Starting invoice field extraction", "INFO")
    logger.log(f"Brand data provided: {bool(brand_data)}", "DEBUG")
    logger.log(f"Message text: '{message_text[:200]}...'", "DEBUG")
    
    values = {}
    
    # Extract invoice number from message
    invoice_number = extract_invoice_number(message_text, logger)
    values["invoice_number"] = invoice_number
    
    # Extract deposit amount from message
    deposit_amount = extract_deposit_amount(message_text, logger)
    values["deposit_amount"] = deposit_amount
    
    # If brand data is provided, use it
    if brand_data:
        logger.log(f"Using brand data: {json.dumps(brand_data, indent=2)}", "DEBUG")
        brand_name = brand_data.get("company_name", "")
        values["brand_name"] = brand_name
        logger.log(f"Brand name from data: {brand_name}", "DEBUG")
        
        # Parse address into components
        address = brand_data.get("address", "")
        address_components = parse_address_components(address, logger)
        
        # Add phone and email if available
        phone = brand_data.get("phone", "Not Available")
        email = brand_data.get("email", "Not Available")
        logger.log(f"Phone: {phone}, Email: {email}", "DEBUG")
    else:
        logger.log("No brand data provided, attempting to extract from message", "WARNING")
        # Try to extract brand name from message
        brand_patterns = [
            r'invoice for\s+([A-Za-z0-9\s&]+?)(?:\s+\d|$)',
            r'generate invoice for\s+([A-Za-z0-9\s&]+?)(?:\s+\d|$)',
        ]
        
        brand_name = ""
        for pattern in brand_patterns:
            match = re.search(pattern, message_text, re.IGNORECASE)
            if match:
                brand_name = match.group(1).strip()
                logger.log(f"Extracted brand name from message: {brand_name}", "DEBUG")
                break
        
        values["brand_name"] = brand_name
        
        # Empty address components
        address_components = parse_address_components("", logger)
        phone = "Not Available"
        email = "Not Available"
    
    # Map to template placeholders (with correct capitalization)
    values["Invoice_Number"] = invoice_number
    values["Brand_Name"] = brand_name
    values["Brand_Address_Line_1"] = address_components["Brand_Address_Line_1"]
    values["Brand_Address_Line_2"] = address_components["Brand_Address_Line_2"]
    values["City"] = address_components["City"]
    values["State"] = address_components["State"]
    values["Pin_Code"] = address_components["Pin_Code"]
    values["Phone"] = phone
    values["Email"] = email
    
    # Calculate dates
    invoice_date = date.today()
    due_date = invoice_date + timedelta(days=15)
    
    values["Invoice_Date"] = invoice_date.strftime("%d/%m/%Y")
    values["Due_Date"] = due_date.strftime("%d/%m/%Y")
    logger.log(f"Invoice Date: {values['Invoice_Date']}, Due Date: {values['Due_Date']}", "DEBUG")
    
    # All three amounts are the same
    if deposit_amount:
        formatted_amount = format_currency(deposit_amount)
        values["Amount_Due"] = formatted_amount
        values["Deposit_Amount"] = formatted_amount
        values["Sub_Total"] = formatted_amount
        logger.log(f"Formatted amount: {formatted_amount}", "SUCCESS")
    else:
        values["Amount_Due"] = ""
        values["Deposit_Amount"] = ""
        values["Sub_Total"] = ""
        logger.log("No deposit amount to format", "WARNING")
    
    # Check for missing required fields
    missing = []
    for field in REQUIRED_FIELDS:
        if not values.get(field):
            missing.append(field)
            logger.log(f"Missing required field: {field}", "WARNING")
    
    logger.log(f"Extraction complete. Extracted {len(values)} fields, {len(missing)} missing", "INFO")
    
    return values, missing


def validate_template_placeholders(doc: Document, logger: InvoiceLogger) -> List[str]:
    """
    Validate which placeholders exist in the template.
    Returns list of found placeholders.
    """
    logger.set_stage("VALIDATE_TEMPLATE")
    logger.log("Scanning template for placeholders", "DEBUG")
    
    placeholders = set()
    
    # Check paragraphs
    for p in doc.paragraphs:
        text = p.text
        matches = re.findall(r'{{([^}]+)}}', text)
        for match in matches:
            placeholders.add(match)
            logger.log(f"Found placeholder in paragraph: {{{{{match}}}}}", "DEBUG")
    
    # Check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text
                    matches = re.findall(r'{{([^}]+)}}', text)
                    for match in matches:
                        placeholders.add(match)
                        logger.log(f"Found placeholder in table: {{{{{match}}}}}", "DEBUG")
    
    logger.log(f"Template validation complete. Found {len(placeholders)} unique placeholders", "SUCCESS")
    return sorted(list(placeholders))


def replace_text_in_paragraph(paragraph, search: str, replace: str, logger: InvoiceLogger) -> bool:
    """
    Helper function to replace text in a paragraph, handling runs correctly.
    Returns True if replacement was made, False otherwise.
    """
    if search not in paragraph.text:
        return False
    
    logger.log(f"Replacing '{search}' with '{replace}' in paragraph", "DEBUG")
    
    # Get the full text
    full_text = paragraph.text
    
    # Check if replacement is actually needed
    if search not in full_text:
        logger.log(f"Search text '{search}' not found in paragraph text", "DEBUG")
        return False
    
    # Replace the search text
    new_text = full_text.replace(search, replace)
    
    # Clear all existing runs except the first one
    while len(paragraph.runs) > 1:
        paragraph._element.remove(paragraph.runs[-1]._element)
    
    # Set the new text on the first run (or add a run if none exist)
    if len(paragraph.runs) > 0:
        paragraph.runs[0].text = new_text
    else:
        paragraph.add_run(new_text)
    
    logger.log(f"Successfully replaced '{search}' with '{replace}'", "SUCCESS")
    return True


def fill_invoice_template(values: dict, output_path: str, logger: InvoiceLogger) -> Dict[str, int]:
    """
    Open the invoice template, replace placeholders, and save to output_path.
    Returns dict with replacement statistics.
    """
    import html
    
    logger.set_stage("FILL_TEMPLATE")
    logger.log(f"Opening template: {TEMPLATE_PATH}", "INFO")
    
    try:
        doc = Document(TEMPLATE_PATH)
    except Exception as e:
        logger.log(f"Failed to open template: {e}", "ERROR")
        raise
    
    # First, validate what placeholders exist in the template
    template_placeholders = validate_template_placeholders(doc, logger)
    logger.log(f"Template contains placeholders: {template_placeholders}", "INFO")
    
    # Track replacements
    replacements = {
        "paragraphs": 0,
        "tables": 0,
        "total": 0
    }
    
    logger.log("Starting placeholder replacement in paragraphs", "INFO")
    
    # Replace in paragraphs
    for i, p in enumerate(doc.paragraphs):
        # Special handling for invoice number with # prefix
        if "#{{Invoice_Number}}" in p.text:
            invoice_value = html.unescape(str(values.get("Invoice_Number", "")))
            if replace_text_in_paragraph(p, "#{{Invoice_Number}}", invoice_value, logger):
                replacements["paragraphs"] += 1
                replacements["total"] += 1
        
        # Replace other placeholders
        for key, val in values.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in p.text:
                clean_value = html.unescape(str(val))
                if replace_text_in_paragraph(p, placeholder, clean_value, logger):
                    replacements["paragraphs"] += 1
                    replacements["total"] += 1
    
    logger.log("Starting placeholder replacement in tables", "INFO")
    
    # Replace in tables
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for paragraph in cell.paragraphs:
                    # Special handling for invoice number with # prefix
                    if "#{{Invoice_Number}}" in paragraph.text:
                        invoice_value = html.unescape(str(values.get("Invoice_Number", "")))
                        if replace_text_in_paragraph(paragraph, "#{{Invoice_Number}}", invoice_value, logger):
                            replacements["tables"] += 1
                            replacements["total"] += 1
                            logger.log(f"Replaced in table {table_idx}, row {row_idx}, cell {cell_idx}", "DEBUG")
                    
                    # Replace other placeholders
                    for key, val in values.items():
                        placeholder = f"{{{{{key}}}}}"
                        if placeholder in paragraph.text:
                            clean_value = html.unescape(str(val))
                            if replace_text_in_paragraph(paragraph, placeholder, clean_value, logger):
                                replacements["tables"] += 1
                                replacements["total"] += 1
                                logger.log(f"Replaced {placeholder} in table {table_idx}, row {row_idx}, cell {cell_idx}", "DEBUG")
    
    logger.log(f"Replacement complete: {replacements['total']} total replacements ({replacements['paragraphs']} in paragraphs, {replacements['tables']} in tables)", "SUCCESS")
    
    # Save document
    try:
        doc.save(output_path)
        logger.log(f"Invoice template saved to: {output_path}", "SUCCESS")
    except Exception as e:
        logger.log(f"Failed to save document: {e}", "ERROR")
        raise
    
    # Verify replacements by checking for remaining placeholders
    verify_doc = Document(output_path)
    remaining_placeholders = validate_template_placeholders(verify_doc, logger)
    
    if remaining_placeholders:
        logger.log(f"WARNING: Found {len(remaining_placeholders)} unreplaced placeholders: {remaining_placeholders}", "WARNING")
    else:
        logger.log("Verification complete: All placeholders replaced successfully", "SUCCESS")
    
    return replacements


def handle_deposit_invoice(event, say, brand_data: Optional[dict] = None):
    """
    Generate deposit invoice with brand data and user-provided amount.
    Multi-step flow with state management.
    
    Args:
        event: Slack event
        say: Slack say function
        brand_data: Optional dict with company_name and address from brand lookup
    """
    raw = event["text"]
    thread_ts = event.get("thread_ts") or event["ts"]
    channel = event["channel"]
    
    # Initialize logger
    logger = InvoiceLogger(thread_ts)
    logger.set_stage("INITIALIZATION")
    logger.log("Starting deposit invoice generation", "INFO")
    logger.log(f"Thread TS: {thread_ts}", "DEBUG")
    logger.log(f"Channel: {channel}", "DEBUG")
    logger.log(f"Brand data provided: {bool(brand_data)}", "INFO")
    
    # Check if we're in an existing flow
    existing_state = get_deposit_invoice_state(thread_ts)
    logger.log(f"Existing state: {existing_state}", "DEBUG")
    
    # Clean text
    logger.set_stage("TEXT_CLEANING")
    cleaned = clean_text(raw)
    logger.log(f"Cleaned text: '{cleaned[:200]}...'", "DEBUG")
    
    # Handle multi-step flow based on state
    if existing_state:
        stage = existing_state.get("stage")
        logger.log(f"Continuing flow at stage: {stage}", "INFO")
        
        if stage == "awaiting_amount":
            # Extract amount from current message
            amount = extract_deposit_amount(cleaned, logger)
            if amount:
                # Store amount and move to next stage
                stored_brand_data = existing_state.get("brand_data")
                set_deposit_invoice_state(thread_ts, "awaiting_invoice_number", stored_brand_data, amount)
                say("💰 Great! Now please provide the invoice number (e.g., 'INV-001' or 'SB/DP/001')", thread_ts=thread_ts)
                return True
            else:
                say("❌ I couldn't find an amount in your message. Please provide it as a number (e.g., '50000' or 'Rs 50000')", thread_ts=thread_ts)
                return False
        
        elif stage == "awaiting_invoice_number":
            # Extract invoice number from current message
            invoice_number = extract_invoice_number(cleaned, logger)
            if invoice_number:
                # We have everything - generate the invoice
                stored_brand_data = existing_state.get("brand_data")
                stored_amount = existing_state.get("amount")
                
                # Create a combined message with all info for extraction
                combined_message = f"invoice {invoice_number} for {stored_amount}"
                logger.log(f"Combined message for extraction: {combined_message}", "DEBUG")
                
                values, missing = extract_invoice_fields(combined_message, stored_brand_data, logger)
                
                if not missing:
                    # Clear state before generating
                    clear_deposit_invoice_state(thread_ts)
                    # Continue with invoice generation below
                else:
                    logger.log(f"Still missing fields after combining: {missing}", "ERROR")
                    clear_deposit_invoice_state(thread_ts)
                    say(f"❌ Error: Missing required fields: {', '.join(missing)}", thread_ts=thread_ts)
                    return False
            else:
                say("❌ I couldn't find an invoice number in your message. Please provide it (e.g., 'INV-001', '#123', or 'SB/DP/001')", thread_ts=thread_ts)
                return False
    else:
        # First time - check what we have
        logger.log("Starting new deposit invoice flow", "INFO")
        values, missing = extract_invoice_fields(cleaned, brand_data, logger)
        
        # If we have brand data but missing amount/invoice, start interactive flow
        if brand_data and missing:
            logger.log("Brand data present but missing other fields, starting interactive flow", "INFO")
            
            if "deposit_amount" in missing and "invoice_number" in missing:
                # Ask for amount first
                set_deposit_invoice_state(thread_ts, "awaiting_amount", brand_data)
                say(f"💰 Great! I have the brand details for **{brand_data.get('company_name', 'the brand')}**.\n\nNow, please provide the deposit amount (e.g., '50000' or 'Rs 50000')", thread_ts=thread_ts)
                return True
            elif "deposit_amount" in missing:
                # Have invoice but need amount
                set_deposit_invoice_state(thread_ts, "awaiting_amount", brand_data)
                say("💰 Please provide the deposit amount (e.g., '50000' or 'Rs 50000')", thread_ts=thread_ts)
                return True
            elif "invoice_number" in missing:
                # Have amount but need invoice
                amount = extract_deposit_amount(cleaned, logger)
                set_deposit_invoice_state(thread_ts, "awaiting_invoice_number", brand_data, amount)
                say("💰 Please provide the invoice number (e.g., 'INV-001' or 'SB/DP/001')", thread_ts=thread_ts)
                return True
    
    # Check for missing fields (after state checks)
    if missing:
        logger.log(f"Missing required fields: {missing}", "ERROR")
        missing_display = {
            "brand_name": "brand name",
            "deposit_amount": "deposit amount (e.g., '5000' or 'Rs 5000')",
            "invoice_number": "invoice number (e.g., 'INV-001' or '#123')"
        }
        missing_names = [missing_display.get(f, f) for f in missing]
        say(f"💰 Please provide the following:\n• " + "\n• ".join(missing_names), thread_ts=thread_ts)
        
        # Send log summary
        say(f"```\n{logger.get_summary()}\n```", thread_ts=thread_ts)
        return False
    
    logger.log("All required fields extracted successfully", "SUCCESS")
    
    # Generate filename
    brand_slug = re.sub(r"\W+", "_", values["brand_name"].lower())
    docx_path = f"{brand_slug}_deposit_invoice.docx"
    pdf_path = f"{brand_slug}_deposit_invoice.pdf"
    logger.log(f"Generated filenames: DOCX={docx_path}, PDF={pdf_path}", "DEBUG")
    
    # Fill template
    try:
        replacements = fill_invoice_template(values, docx_path, logger)
        logger.log(f"Template filled successfully with {replacements['total']} replacements", "SUCCESS")
    except Exception as e:
        logger.log(f"Error filling invoice template: {e}", "ERROR")
        say(f"❌ Sorry, I encountered an error generating the invoice: {e}", thread_ts=thread_ts)
        say(f"```\n{logger.get_summary()}\n```", thread_ts=thread_ts)
        return False
    
    # Convert to PDF
    logger.set_stage("PDF_CONVERSION")
    pdf_ok = convert_docx_to_pdf(docx_path, pdf_path)
    
    if pdf_ok:
        logger.log("PDF conversion successful", "SUCCESS")
        upload_path = pdf_path
        title = f"{values['Brand_Name']} Deposit Invoice (PDF)"
    else:
        logger.log("PDF conversion failed, will upload DOCX", "WARNING")
        upload_path = docx_path
        title = f"{values['Brand_Name']} Deposit Invoice (Word)"
    
    # Upload to Slack
    logger.set_stage("SLACK_UPLOAD")
    try:
        with open(upload_path, "rb") as f:
            slack_client.files_upload_v2(
                channel=channel,
                thread_ts=thread_ts,
                file=f,
                filename=os.path.basename(upload_path),
                title=title,
                initial_comment=f"📎 Here's your *{title}*"
            )
        logger.log(f"Invoice uploaded successfully: {title}", "SUCCESS")
        
        # Send success summary
        summary = f"✅ *Invoice Generation Complete!*\n\n"
        summary += f"• Brand: {values['Brand_Name']}\n"
        summary += f"• Invoice #: {values['Invoice_Number']}\n"
        summary += f"• Amount: {values['Deposit_Amount']}\n"
        summary += f"• Replacements Made: {replacements['total']}\n"
        summary += f"• Format: {'PDF' if pdf_ok else 'DOCX'}"
        say(summary, thread_ts=thread_ts)
        
        # Send detailed log in a code block for debugging
        say(f"📋 *Debug Log:*\n```\n{logger.get_summary()}\n```", thread_ts=thread_ts)
        
        return True
    except Exception as e:
        logger.log(f"Error uploading invoice: {e}", "ERROR")
        say(f"❌ Invoice generated but upload failed: {e}", thread_ts=thread_ts)
        say(f"```\n{logger.get_summary()}\n```", thread_ts=thread_ts)
        return False


# Test function
if __name__ == "__main__":
    print("🧪 Testing Deposit Invoice Service V2")
    print("=" * 50)
    
    # Create a test logger
    test_logger = InvoiceLogger("test-thread")
    
    # Test amount extraction
    test_messages = [
        "generate invoice for 50000",
        "deposit invoice Rs 10000",
        "create invoice amount 7500",
        "₹15000 invoice",
        "invoice SB/DP/001 for 50000"
    ]
    
    print("\n📝 Testing Amount Extraction:")
    for msg in test_messages:
        print(f"\n  Message: {msg}")
        amount = extract_deposit_amount(msg, test_logger)
        print(f"  Result: {amount}")
    
    # Test invoice number extraction
    print("\n📝 Testing Invoice Number Extraction:")
    for msg in test_messages:
        print(f"\n  Message: {msg}")
        invoice_num = extract_invoice_number(msg, test_logger)
        print(f"  Result: {invoice_num}")
    
    # Test field extraction with brand data
    print("\n" + "=" * 50)
    print("📝 Testing Field Extraction with Brand Data:")
    
    test_brand_data = {
        "company_name": "Test Company",
        "address": "123 Test Street, Test Area, Mumbai, Maharashtra, 400001",
        "phone": "+91 9876543210",
        "email": "test@company.com"
    }
    
    values, missing = extract_invoice_fields(
        "generate invoice SB/DP/001 for 50000", 
        test_brand_data, 
        test_logger
    )
    
    print(f"\n  Extracted Values:")
    for key, val in values.items():
        print(f"    {key}: {val}")
    print(f"\n  Missing Fields: {missing}")
    
    print("\n" + "=" * 50)
    print("📊 Logger Summary:")
    print(test_logger.get_summary())
