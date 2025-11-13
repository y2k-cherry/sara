#!/usr/bin/env python3
"""
Deposit Invoice Generator Service
Generates deposit invoices using the Advance Deposit Invoice Template
"""

import os
import re
from datetime import date, timedelta
from docx import Document
from slack_sdk import WebClient
from dotenv import load_dotenv

# Google Docs-based PDF exporter
from google_pdf import convert_docx_to_pdf_google as convert_docx_to_pdf

load_dotenv()

# Constants
SLACK_TOKEN = os.getenv("SLACK_BOT_TOKEN")
TEMPLATE_PATH = "Advance Deposit Invoice Template.docx"

# Initialize Slack client
slack_client = WebClient(token=SLACK_TOKEN)

# Required fields for invoice generation
REQUIRED_FIELDS = [
    "brand_name",
    "deposit_amount",
    "invoice_number"
]


def clean_text(text: str) -> str:
    """Strip Slack mentions and extra whitespace."""
    return re.sub(r"<@[\w]+>", "", text).strip()


def format_currency(value: str) -> str:
    """Format an integer/float string as Indian-rupee currency."""
    try:
        num = float(value)
        return f"₹{num:,.0f}"
    except:
        return value


def convert_number_to_words(number_str: str) -> str:
    """Convert a number string to words (Indian numbering system)."""
    try:
        num = int(float(number_str))
        
        if num == 0:
            return "zero"
        elif num < 0:
            return "negative " + convert_number_to_words(str(-num))
        
        # Indian numbering system
        ones = ["", "one", "two", "three", "four", "five", "six", "seven", "eight", "nine"]
        teens = ["ten", "eleven", "twelve", "thirteen", "fourteen", "fifteen", 
                "sixteen", "seventeen", "eighteen", "nineteen"]
        tens = ["", "", "twenty", "thirty", "forty", "fifty", "sixty", "seventy", "eighty", "ninety"]
        
        def convert_hundreds(n):
            result = ""
            if n >= 100:
                result += ones[n // 100] + " hundred "
                n %= 100
            if n >= 20:
                result += tens[n // 10] + " "
                n %= 10
            elif n >= 10:
                result += teens[n - 10] + " "
                n = 0
            if n > 0:
                result += ones[n] + " "
            return result.strip()
        
        result = ""
        if num >= 10000000:  # 1 crore
            crores = num // 10000000
            result = convert_hundreds(crores) + " crore "
            num %= 10000000
            
        if num >= 100000:  # 1 lakh
            lakhs = num // 100000
            result += convert_hundreds(lakhs) + " lakh "
            num %= 100000
            
        if num >= 1000:  # 1 thousand
            thousands = num // 1000
            result += convert_hundreds(thousands) + " thousand "
            num %= 1000
            
        if num > 0:
            result += convert_hundreds(num)
            
        return result.strip()
        
    except:
        return number_str


def extract_deposit_amount(message_text: str) -> str:
    """
    Extract deposit amount from the message.
    Patterns: "5000", "Rs 5000", "₹5000", "deposit 5000", "amount 5000" etc.
    """
    # Try various patterns - more specific patterns first
    patterns = [
        r'(?:amount|deposit)\s+(?:of\s+)?(?:rs\.?|₹)?\s*([0-9,]+)',
        r'(?:for|invoice\s+for)\s+(?:rs\.?|₹)?\s*([0-9,]+)',
        r'(?:rs\.?|₹)\s*([0-9,]+)',
        r'\b([0-9]{4,})\b',  # Any 4+ digit number as fallback
    ]
    
    for pattern in patterns:
        match = re.search(pattern, message_text, re.IGNORECASE)
        if match:
            amount = match.group(1).replace(',', '')
            print(f"💰 Extracted deposit amount: {amount}")
            return amount
    
    print("⚠️ No deposit amount found in message")
    return ""


def extract_invoice_number(message_text: str) -> str:
    """
    Extract invoice number from the message.
    Patterns: "invoice #123", "invoice number 123", "#INV-001", etc.
    """
    patterns = [
        r'#\s*([A-Z0-9-]+)',  # Match #INV-001, #123, etc.
        r'invoice\s+#\s*([A-Z0-9-]+)',  # Match "invoice #123"
        r'invoice\s+number\s*:?\s*([A-Z0-9-]+)',  # Match "invoice number 123"
        r'\b(INV-[0-9]+)\b',  # Match standalone INV-001
    ]
    
    for pattern in patterns:
        match = re.search(pattern, message_text, re.IGNORECASE)
        if match:
            invoice_num = match.group(1).strip()
            print(f"📋 Extracted invoice number: {invoice_num}")
            return invoice_num
    
    print("⚠️ No invoice number found in message")
    return ""


def parse_address_components(address: str) -> dict:
    """
    Parse a combined address into separate components.
    Returns dict with Brand_Address_Line_1, Brand_Address_Line_2, City, State, Pin_Code
    """
    if not address:
        return {
            "Brand_Address_Line_1": "",
            "Brand_Address_Line_2": "",
            "City": "",
            "State": "",
            "Pin_Code": ""
        }
    
    # Split address by commas
    parts = [p.strip() for p in address.split(',')]
    
    # Initialize components
    components = {
        "Brand_Address_Line_1": "",
        "Brand_Address_Line_2": "",
        "City": "",
        "State": "",
        "Pin_Code": ""
    }
    
    # Try to extract pin code (6 digits)
    pin_code_pattern = r'\b(\d{6})\b'
    for i, part in enumerate(parts):
        match = re.search(pin_code_pattern, part)
        if match:
            components["Pin_Code"] = match.group(1)
            # Remove pin code from this part
            parts[i] = re.sub(pin_code_pattern, '', part).strip()
            break
    
    # Remove empty parts
    parts = [p for p in parts if p]
    
    # Assign remaining parts
    if len(parts) >= 1:
        components["Brand_Address_Line_1"] = parts[0]
    if len(parts) >= 2:
        components["Brand_Address_Line_2"] = parts[1]
    if len(parts) >= 3:
        components["City"] = parts[2]
    if len(parts) >= 4:
        components["State"] = parts[3]
    
    # If we have fewer parts, try to be smart about city/state
    if len(parts) == 3 and not components["State"]:
        # Assume last part is state if no state found yet
        components["State"] = parts[2]
        components["City"] = parts[1] if len(parts) > 1 else ""
    
    return components


def extract_invoice_fields(message_text: str, brand_data: dict = None):
    """
    Extract invoice fields from message and brand data.
    
    Args:
        message_text: The user's message
        brand_data: Optional brand data from brand_info_service with:
            - company_name
            - address (combined)
            - phone (optional)
            - email (optional)
    
    Returns:
        Tuple of (values dict, list of missing fields)
    """
    print(f"🔍 Starting invoice field extraction...")
    print(f"🔍 Brand data provided: {bool(brand_data)}")
    print(f"🔍 Message: {message_text[:100]}...")
    
    values = {}
    
    # Extract invoice number from message
    invoice_number = extract_invoice_number(message_text)
    values["invoice_number"] = invoice_number
    
    # Extract deposit amount from message
    deposit_amount = extract_deposit_amount(message_text)
    values["deposit_amount"] = deposit_amount
    
    # If brand data is provided, use it
    if brand_data:
        print(f"🔍 Using brand data: {brand_data}")
        brand_name = brand_data.get("company_name", "")
        values["brand_name"] = brand_name
        
        # Parse address into components
        address = brand_data.get("address", "")
        address_components = parse_address_components(address)
        
        # Add phone and email if available
        phone = brand_data.get("phone", "Not Available")
        email = brand_data.get("email", "Not Available")
    else:
        # Try to extract brand name from message
        brand_patterns = [
            r'invoice for\s+([A-Za-z0-9\s&]+?)(?:\s+\d|$)',
            r'generate invoice for\s+([A-Za-z0-9\s&]+?)(?:\s+\d|$)',
        ]
        
        brand_name = ""
        for pattern in brand_patterns:
            match = re.search(pattern, message_text, re.IGNORECASE)
            if match:
                brand_name = match.group(1).strip()
                break
        
        values["brand_name"] = brand_name
        
        # Empty address components
        address_components = parse_address_components("")
        phone = "Not Available"
        email = "Not Available"
    
    # Map to template placeholders (with correct capitalization)
    # NOTE: Invoice_Number in template has a # prefix
    values["Invoice_Number"] = invoice_number
    values["Brand_Name"] = brand_name
    values["Brand_Address_Line_1"] = address_components["Brand_Address_Line_1"]
    values["Brand_Address_Line_2"] = address_components["Brand_Address_Line_2"]
    values["City"] = address_components["City"]
    values["State"] = address_components["State"]
    values["Pin_Code"] = address_components["Pin_Code"]
    values["Phone"] = phone
    values["Email"] = email
    
    # Calculate dates
    invoice_date = date.today()
    due_date = invoice_date + timedelta(days=15)
    
    values["Invoice_Date"] = invoice_date.strftime("%d/%m/%Y")
    values["Due_Date"] = due_date.strftime("%d/%m/%Y")
    
    # All three amounts are the same
    if deposit_amount:
        formatted_amount = format_currency(deposit_amount)
        values["Amount_Due"] = formatted_amount
        values["Deposit_Amount"] = formatted_amount
        values["Sub_Total"] = formatted_amount
    else:
        values["Amount_Due"] = ""
        values["Deposit_Amount"] = ""
        values["Sub_Total"] = ""
    
    # Check for missing required fields
    missing = []
    for field in REQUIRED_FIELDS:
        if not values.get(field):
            missing.append(field)
    
    print(f"🔍 Extracted values: {values}")
    print(f"🔍 Missing fields: {missing}")
    
    return values, missing


def replace_text_in_paragraph(paragraph, search, replace):
    """Helper function to replace text in a paragraph, handling runs correctly."""
    if search not in paragraph.text:
        return
    
    # Get the full text
    full_text = paragraph.text
    
    # Replace the search text
    new_text = full_text.replace(search, replace)
    
    # Clear all existing runs except the first one
    while len(paragraph.runs) > 1:
        paragraph._element.remove(paragraph.runs[-1]._element)
    
    # Set the new text on the first run (or add a run if none exist)
    if len(paragraph.runs) > 0:
        paragraph.runs[0].text = new_text
    else:
        paragraph.add_run(new_text)


def fill_invoice_template(values: dict, output_path: str):
    """
    Open the invoice template, replace placeholders, and save to output_path.
    Uses proper Word document text replacement to handle runs correctly.
    """
    import html
    
    doc = Document(TEMPLATE_PATH)
    
    # Replace in paragraphs
    for p in doc.paragraphs:
        # Check for invoice number with # prefix
        if "#{{Invoice_Number}}" in p.text:
            replace_text_in_paragraph(p, "#{{Invoice_Number}}", 
                                    html.unescape(str(values.get("Invoice_Number", ""))))
        
        # Replace other placeholders
        for key, val in values.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in p.text:
                replace_text_in_paragraph(p, placeholder, html.unescape(str(val)))
    
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    # Check for invoice number with # prefix
                    if "#{{Invoice_Number}}" in paragraph.text:
                        replace_text_in_paragraph(paragraph, "#{{Invoice_Number}}", 
                                                html.unescape(str(values.get("Invoice_Number", ""))))
                    
                    # Replace other placeholders
                    for key, val in values.items():
                        placeholder = f"{{{{{key}}}}}"
                        if placeholder in paragraph.text:
                            replace_text_in_paragraph(paragraph, placeholder, 
                                                    html.unescape(str(val)))
    
    doc.save(output_path)
    print(f"✅ Invoice template filled and saved to: {output_path}")
    print(f"✅ Replaced placeholders with actual values")


def handle_deposit_invoice(event, say, brand_data: dict = None):
    """
    Generate deposit invoice with brand data and user-provided amount.
    
    Args:
        event: Slack event
        say: Slack say function
        brand_data: Optional dict with company_name and address from brand lookup
    """
    raw = event["text"]
    thread_ts = event.get("thread_ts") or event["ts"]
    channel = event["channel"]
    
    print(f"📋 Starting deposit invoice generation...")
    print(f"📋 Brand data provided: {bool(brand_data)}")
    
    # Send acknowledgement
    say("🧾 Got it - working on your deposit invoice! ⚡", thread_ts=thread_ts)
    
    # Clean and extract fields
    cleaned = clean_text(raw)
    values, missing = extract_invoice_fields(cleaned, brand_data)
    
    # Check for missing fields
    if missing:
        missing_display = {
            "brand_name": "brand name",
            "deposit_amount": "deposit amount (e.g., '5000' or 'Rs 5000')",
            "invoice_number": "invoice number (e.g., 'INV-001' or '#123')"
        }
        missing_names = [missing_display.get(f, f) for f in missing]
        say(f"💰 Please provide the following:\n• " + "\n• ".join(missing_names), thread_ts=thread_ts)
        return False
    
    # Generate filename
    brand_slug = re.sub(r"\W+", "_", values["brand_name"].lower())
    docx_path = f"{brand_slug}_deposit_invoice.docx"
    pdf_path = f"{brand_slug}_deposit_invoice.pdf"
    
    # Fill template
    try:
        fill_invoice_template(values, docx_path)
    except Exception as e:
        print(f"❌ Error filling invoice template: {e}")
        say(f"❌ Sorry, I encountered an error generating the invoice: {e}", thread_ts=thread_ts)
        return False
    
    # Convert to PDF
    pdf_ok = convert_docx_to_pdf(docx_path, pdf_path)
    
    # Decide which to upload
    if pdf_ok:
        upload_path = pdf_path
        title = f"{values['brand_name']} Deposit Invoice (PDF)"
    else:
        upload_path = docx_path
        title = f"{values['brand_name']} Deposit Invoice (Word)"
    
    # Upload to Slack
    try:
        with open(upload_path, "rb") as f:
            slack_client.files_upload_v2(
                channel=channel,
                thread_ts=thread_ts,
                file=f,
                filename=os.path.basename(upload_path),
                title=title,
                initial_comment=f"📎 Here's your *{title}*"
            )
        print(f"✅ Invoice uploaded successfully: {title}")
        return True
    except Exception as e:
        print(f"❌ Error uploading invoice: {e}")
        say(f"❌ Invoice generated but upload failed: {e}", thread_ts=thread_ts)
        return False


# Test function
if __name__ == "__main__":
    print("🧪 Testing Deposit Invoice Service")
    print("=" * 50)
    
    # Test amount extraction
    test_messages = [
        "generate invoice for 5000",
        "deposit invoice Rs 10000",
        "create invoice amount 7500",
        "₹15000 invoice",
    ]
    
    for msg in test_messages:
        print(f"\nTest message: {msg}")
        amount = extract_deposit_amount(msg)
        print(f"Extracted amount: {amount}")
    
    # Test field extraction
    print("\n" + "=" * 50)
    print("Testing field extraction with brand data:")
    
    test_brand_data = {
        "company_name": "Test Company",
        "address": "123 Test Street, Test City"
    }
    
    values, missing = extract_invoice_fields("generate invoice for 5000", test_brand_data)
    print(f"Values: {values}")
    print(f"Missing: {missing}")
