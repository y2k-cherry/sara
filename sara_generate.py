import os
from openai import OpenAI
from dotenv import load_dotenv
from docx import Document
from datetime import date

load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))


def extract_values_from_prompt(prompt):
    system_msg = (
        "You are Sara, a helpful sales back-office assistant. "
        "Extract all fields needed to fill a brand partnership agreement. "
        "Return a JSON object with these keys ONLY: "
        "brand_name, company_name, company_address, industry, flat_fee, deposit, deposit_in_words, start_date. "
        "Omit optional text, no extra explanation."
    )

    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": system_msg},
            {"role": "user", "content": prompt}
        ]
    )

    reply = response.choices[0].message.content
    try:
        data = eval(reply)  # safe for trusted input like GPT's structured JSON
        return data
    except Exception as e:
        print("❌ Error parsing response from GPT:\n", reply)
        raise e


def fill_missing_fields(values):
    expected_fields = {
        "brand_name": "Brand Name",
        "company_name": "Company Legal Name",
        "company_address": "Company Address",
        "industry": "Industry",
        "flat_fee": "Flat Fee (e.g. 1000)",
        "deposit": "Deposit Amount (e.g. 10000)",
        "deposit_in_words": "Deposit in Words (e.g. Ten Thousand Only)",
        "start_date": f"Agreement Date [default: today {date.today()}]"
    }

    for key, label in expected_fields.items():
        if key not in values or not values[key]:
            user_input = input(f"{label}: ")
            values[key] = user_input if user_input else str(date.today()) if key == "start_date" else ""

    return values


def generate_agreement(input_path, output_path, values):
    doc = Document(input_path)

    for para in doc.paragraphs:
        for key, val in values.items():
            if f"{{{{{key}}}}}" in para.text:
                para.text = para.text.replace(f"{{{{{key}}}}}", val)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, val in values.items():
                    if f"{{{{{key}}}}}" in cell.text:
                        cell.text = cell.text.replace(f"{{{{{key}}}}}", val)

    doc.save(output_path)
    print(f"\n✅ Agreement saved as: {output_path}")


def main():
    print("\n💬 Welcome to Sara — your sales ops AI agent\n")

    prompt = input("🗣 What would you like Sara to do?\n> ")

    print("\n🤖 Extracting details from your request...")
    extracted_values = extract_values_from_prompt(prompt)

    print("🔍 Checking for missing fields...")
    filled_values = fill_missing_fields(extracted_values)

    file_name = f"{filled_values['brand_name'].replace(' ', '_')}_Agreement.docx"
    generate_agreement("Partnership Agreement Template.docx", file_name, filled_values)


if __name__ == "__main__":
    main()
